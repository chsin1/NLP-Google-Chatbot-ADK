#!/usr/bin/env python3
"""Sync DOCX files from same-named Markdown files in a docs directory."""

from __future__ import annotations

import argparse
from pathlib import Path

import markdown
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.text.paragraph import Paragraph


def _append_inline(paragraph: Paragraph, node, *, bold: bool = False, italic: bool = False, monospace: bool = False) -> None:
    if isinstance(node, NavigableString):
        text = str(node)
        if not text:
            return
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        if monospace:
            run.font.name = "Courier New"
        return

    if not isinstance(node, Tag):
        return

    tag = node.name.lower()
    if tag in {"strong", "b"}:
        for child in node.children:
            _append_inline(paragraph, child, bold=True or bold, italic=italic, monospace=monospace)
        return
    if tag in {"em", "i"}:
        for child in node.children:
            _append_inline(paragraph, child, bold=bold, italic=True or italic, monospace=monospace)
        return
    if tag == "code":
        for child in node.children:
            _append_inline(paragraph, child, bold=bold, italic=italic, monospace=True)
        return
    if tag == "a":
        label = node.get_text(" ", strip=True)
        href = node.attrs.get("href", "")
        text = f"{label} ({href})" if href else label
        if text:
            run = paragraph.add_run(text)
            run.bold = bold
            run.italic = italic
        return
    if tag == "br":
        paragraph.add_run().add_break()
        return

    for child in node.children:
        _append_inline(paragraph, child, bold=bold, italic=italic, monospace=monospace)


def _render_list(doc: Document, list_node: Tag, *, ordered: bool, depth: int = 0) -> None:
    style_name = "List Number" if ordered else "List Bullet"
    for li in list_node.find_all("li", recursive=False):
        paragraph = doc.add_paragraph(style=style_name)
        if depth > 0:
            # Indent nested list items to keep hierarchy readable.
            paragraph.paragraph_format.left_indent = paragraph.part.document.styles["Normal"].paragraph_format.left_indent
            if paragraph.paragraph_format.left_indent is None:
                from docx.shared import Inches

                paragraph.paragraph_format.left_indent = Inches(0.25 * depth)
        for child in li.children:
            if isinstance(child, Tag) and child.name.lower() in {"ul", "ol"}:
                continue
            _append_inline(paragraph, child)
        for nested in li.find_all(["ul", "ol"], recursive=False):
            _render_list(doc, nested, ordered=nested.name.lower() == "ol", depth=depth + 1)


def _render_table(doc: Document, table_node: Tag) -> None:
    rows = table_node.find_all("tr")
    if not rows:
        return
    max_cols = max(len(row.find_all(["th", "td"])) for row in rows)
    if max_cols == 0:
        return

    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"

    for row_idx, row in enumerate(rows):
        cells = row.find_all(["th", "td"])
        for col_idx, cell in enumerate(cells):
            target = table.cell(row_idx, col_idx)
            target.text = ""
            paragraph = target.paragraphs[0]
            for child in cell.children:
                _append_inline(paragraph, child)
            if cell.name.lower() == "th":
                for run in paragraph.runs:
                    run.bold = True


def _render_block(doc: Document, node) -> None:
    if isinstance(node, NavigableString):
        text = str(node).strip()
        if text:
            doc.add_paragraph(text)
        return

    if not isinstance(node, Tag):
        return

    tag = node.name.lower()
    if tag in {"h1", "h2", "h3", "h4", "h5", "h6"}:
        level = int(tag[1])
        paragraph = doc.add_heading(level=min(level, 4))
        for child in node.children:
            _append_inline(paragraph, child)
        return
    if tag == "p":
        paragraph = doc.add_paragraph()
        for child in node.children:
            _append_inline(paragraph, child)
        return
    if tag == "ul":
        _render_list(doc, node, ordered=False)
        return
    if tag == "ol":
        _render_list(doc, node, ordered=True)
        return
    if tag == "pre":
        text = node.get_text("\n", strip=False)
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        run.font.name = "Courier New"
        return
    if tag == "table":
        _render_table(doc, node)
        return
    if tag == "blockquote":
        paragraph = doc.add_paragraph()
        for child in node.children:
            _append_inline(paragraph, child, italic=True)
        return
    if tag == "hr":
        doc.add_paragraph("")
        return

    for child in node.children:
        _render_block(doc, child)


def markdown_to_docx(md_path: Path, docx_path: Path) -> None:
    markdown_text = md_path.read_text(encoding="utf-8")
    html = markdown.markdown(
        markdown_text,
        extensions=["fenced_code", "tables", "sane_lists"],
        output_format="html5",
    )
    soup = BeautifulSoup(html, "html.parser")
    doc = Document()

    for child in soup.children:
        _render_block(doc, child)

    doc.save(docx_path)


def sync_docs(docs_dir: Path) -> list[tuple[Path, Path]]:
    synced: list[tuple[Path, Path]] = []
    for md_path in sorted(docs_dir.glob("*.md")):
        docx_path = md_path.with_suffix(".docx")
        if not docx_path.exists():
            continue
        markdown_to_docx(md_path, docx_path)
        synced.append((md_path, docx_path))
    return synced


def main() -> int:
    parser = argparse.ArgumentParser(description="Update .docx files from same-name .md files.")
    parser.add_argument("--docs-dir", default="docs", help="Path to docs directory (default: docs)")
    args = parser.parse_args()

    docs_dir = Path(args.docs_dir).resolve()
    if not docs_dir.exists():
        raise SystemExit(f"Docs directory does not exist: {docs_dir}")

    synced = sync_docs(docs_dir)
    if not synced:
        print(f"No matching markdown/docx pairs found in {docs_dir}")
        return 0

    print("Updated DOCX files:")
    for md_path, docx_path in synced:
        print(f"- {docx_path.name} <- {md_path.name}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
